# webCollect.py

from playwright.sync_api import sync_playwright, TimeoutError
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, RGBColor

import requests
import os
import re
import time
import uuid
import shutil


# ─────────────────────────────────────────────
# PATH
# ─────────────────────────────────────────────

SAVE_DIR = r"C:\KHK\DocuWork\webdocu"
IMG_FOLDER = os.path.join(SAVE_DIR, "tmp_imgs")

os.makedirs(SAVE_DIR, exist_ok=True)
os.makedirs(IMG_FOLDER, exist_ok=True)


# ─────────────────────────────────────────────
# IMAGE DOWNLOAD
# ─────────────────────────────────────────────

def download_image(url, cookies=None):

    try:
        headers = {
            "Referer": "https://memo.naver.com"
        }

        r = requests.get(
            url,
            headers=headers,
            cookies=cookies,
            timeout=15
        )

        if r.status_code != 200:
            return None

        if "image" not in r.headers.get("Content-Type", ""):
            return None

        ext = ".png"

        content_type = r.headers.get("Content-Type", "")

        if "jpeg" in content_type:
            ext = ".jpg"
        elif "gif" in content_type:
            ext = ".gif"
        elif "webp" in content_type:
            ext = ".webp"

        filename = f"{uuid.uuid4().hex}{ext}"

        path = os.path.join(IMG_FOLDER, filename)

        with open(path, "wb") as f:
            f.write(r.content)

        return path

    except Exception as e:
        print(f"    ⚠️ 이미지 다운로드 실패: {e}")
        return None


# ─────────────────────────────────────────────
# INLINE TEXT PROCESS
# ─────────────────────────────────────────────

def apply_style(run, tag_name):

    if tag_name in ["b", "strong"]:
        run.bold = True

    if tag_name in ["i", "em"]:
        run.italic = True

    if tag_name == "u":
        run.underline = True

    if tag_name in ["s", "strike", "del"]:
        run.font.strike = True


def process_inline(node, para):

    if isinstance(node, NavigableString):

        text = str(node).replace("\xa0", " ")

        if text.strip():
            para.add_run(text)

        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        para.add_run("\n")
        return

    for child in node.children:

        if isinstance(child, NavigableString):

            text = str(child).replace("\xa0", " ")

            if not text.strip():
                continue

            run = para.add_run(text)

            parents = [p.name for p in child.parents]

            for p in parents:
                apply_style(run, p)

        else:
            process_inline(child, para)


# ─────────────────────────────────────────────
# BLOCK PROCESS
# ─────────────────────────────────────────────

def process_block(block, doc, cookies=None):

    para = doc.add_paragraph()

    for child in block.children:

        if isinstance(child, Tag) and child.name == "img":

            src = child.get("src", "")

            if src.startswith("http"):

                img_path = download_image(src, cookies)

                if img_path:
                    try:
                        doc.add_picture(img_path, width=Inches(4.8))
                    except:
                        pass

            continue

        process_inline(child, para)


# ─────────────────────────────────────────────
# MEMO HTML → DOCX
# ─────────────────────────────────────────────

def memo_to_doc(html, doc, index, cookies=None):

    soup = BeautifulSoup(html, "html.parser")

    folder = soup.select_one(".folder_name .text")
    stamp = soup.select_one(".time_stamp span")
    title_el = soup.select_one("h2.edit_title")
    content = soup.select_one(".workseditor-content")

    folder_text = folder.get_text(strip=True) if folder else ""
    stamp_text = stamp.get_text(strip=True) if stamp else ""
    title_text = title_el.get_text(strip=True) if title_el else f"메모 {index}"

    doc.add_heading(title_text, level=2)

    meta = doc.add_paragraph()

    run = meta.add_run(
        f"📁 {folder_text}   |   🕒 {stamp_text}"
    )

    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    if content:

        for child in content.children:

            if not isinstance(child, Tag):
                continue

            if child.name in ["div", "p", "li"]:
                process_block(child, doc, cookies)

            elif child.name in ["ul", "ol"]:

                for li in child.find_all("li", recursive=False):
                    process_block(li, doc, cookies)

    doc.add_paragraph("━" * 35)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# FOLDER LIST
# ─────────────────────────────────────────────

def get_folder_tabs(page):

    page.wait_for_selector("div.memo")

    result = []

    tabs = page.query_selector_all("div.memo")

    for tab in tabs:

        try:
            name_el = tab.query_selector("span.text")
            count_el = tab.query_selector("span.count")

            if not name_el:
                continue

            name = name_el.inner_text().strip()

            count = "0"

            if count_el:
                count = count_el.inner_text().strip()

            if name[:4].isdigit():

                result.append({
                    "name": name,
                    "count": int(count.replace(",", ""))
                })

        except:
            continue

    return result


# ─────────────────────────────────────────────
# CLICK FOLDER
# ─────────────────────────────────────────────

def click_folder(page, folder_name):

    tabs = page.query_selector_all("div.memo")

    for tab in tabs:

        try:
            name_el = tab.query_selector("span.text")

            if not name_el:
                continue

            name = name_el.inner_text().strip()

            if name == folder_name:

                tab.scroll_into_view_if_needed()

                try:
                    tab.click(timeout=3000)
                except:
                    tab.evaluate("el => el.click()")

                return True

        except:
            continue

    return False


# ─────────────────────────────────────────────
# GET NEXT MEMO
# stale-free 핵심 구조
# ─────────────────────────────────────────────

def get_next_memo(page, processed):

    buttons = page.query_selector_all("button.paper_link")

    for idx, btn in enumerate(buttons):

        try:

            text = btn.inner_text().strip()

            key = f"{idx}_{text[:80]}"

            if key in processed:
                continue

            processed.add(key)

            return btn

        except:
            continue

    return None


# ─────────────────────────────────────────────
# SCROLL
# ─────────────────────────────────────────────

def scroll_memo_list(page):

    page.evaluate("""
        () => {

            const candidates = [
                document.querySelector('.memo_list'),
                document.querySelector('.ReactVirtualized__Grid'),
                document.querySelector('.list_wrap')
            ];

            for (const el of candidates) {

                if (!el) continue;

                el.scrollTop = el.scrollHeight;
            }

            window.scrollBy(0, 3000);
        }
    """)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():

    with sync_playwright() as p:

        browser = p.chromium.launch(
            headless=False,
            args=[ "--start-maximized", 
                "--disable-blink-features=AutomationControlled" 
            ]
        )

        context = browser.new_context(
            viewport={
                "width": 1400,
                "height": 900
            },
            user_agent=( 
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) " 
                "AppleWebKit/537.36 (KHTML, like Gecko) " 
                "Chrome/124.0.0.0 Safari/537.36" 
            )
        )

        page = context.new_page()

        
# LOGIN
        page.goto("https://nid.naver.com/nidlogin.login")

        print("\n✅ 브라우저에서 네이버 로그인 진행하세요.")

        # 로그인 완료 대기
        while True:

            current_url = page.url

            # 로그인 페이지 벗어나면 성공으로 간주
            if "nidlogin" not in current_url:
                break

            time.sleep(1)

        print("  ✅ 로그인 확인 완료")

        # 네이버 메모 이동
        page.goto(
            "https://memo.naver.com",
            wait_until="domcontentloaded"
        )

        page.wait_for_load_state("networkidle")

        # FOLDERS
        folders = get_folder_tabs(page)

        print(f"\n📁 폴더 {len(folders)}개 발견")

        for folder in folders:

            folder_name = folder["name"]
            target_total = folder["count"]

            print("\n" + "=" * 50)
            print(f"📂 {folder_name} ({target_total}개)")
            print("=" * 50)

            safe_name = re.sub(
                r'[\\/*?:"<>|]',
                "",
                folder_name
            )

            save_file = os.path.join(
                SAVE_DIR,
                f"naver_memo_{safe_name}.docx"
            )

            doc = Document()

            doc.add_heading(
                f"네이버 메모 - {folder_name}",
                level=1
            )

            if not click_folder(page, folder_name):

                print("  ⚠️ 폴더 클릭 실패")
                continue

            page.wait_for_timeout(1500)

            # REQUEST COOKIE
            pw_cookies = context.cookies()

            req_cookies = {
                c["name"]: c["value"]
                for c in pw_cookies
            }

            processed = set()

            folder_count = 0
            no_new_count = 0

            while folder_count < target_total:

                next_btn = get_next_memo(
                    page,
                    processed
                )

                # NEW ITEM NOT FOUND
                if not next_btn:

                    no_new_count += 1

                    if no_new_count >= 5:
                        print("  ⚠️ 새 메모 없음")
                        break

                    print(
                        f"  ⬇️ 스크롤 로딩..."
                    )

                    scroll_memo_list(page)

                    page.wait_for_timeout(2000)

                    continue

                no_new_count = 0

                try:

                    next_btn.scroll_into_view_if_needed()

                    try:
                        next_btn.click(timeout=3000)
                    except:
                        next_btn.evaluate(
                            "el => el.click()"
                        )

                    page.wait_for_selector(
                        "div.wrap_app",
                        timeout=5000
                    )

                    wrap = page.query_selector(
                        "div.wrap_app"
                    )

                    if not wrap:
                        raise Exception(
                            "상세 화면 없음"
                        )

                    html = wrap.evaluate(
                        "el => el.outerHTML"
                    )

                    folder_count += 1

                    print(
                        f"  [{folder_count}/{target_total}] 수집 완료"
                    )

                    memo_to_doc(
                        html,
                        doc,
                        folder_count,
                        req_cookies
                    )

                except Exception as e:

                    print(f"    ⚠️ 실패: {e}")

                finally:

                    try:

                        back_btn = page.query_selector(
                            "button.btn_back"
                        )

                        if back_btn:
                            back_btn.click()
                        else:
                            page.go_back()

                    except:
                        pass

                    page.wait_for_timeout(1200)

            # SAVE
            try:

                doc.save(save_file)

                print(f"  💾 저장 완료")
                print(f"     {save_file}")

            except Exception as e:

                print(f"  ⚠️ 저장 실패: {e}")

        browser.close()

    # CLEANUP
    try:
        shutil.rmtree(
            IMG_FOLDER,
            ignore_errors=True
        )
    except:
        pass

    print("\n🎉 전체 완료")


if __name__ == "__main__": 
    try: 
        main() 
    except Exception as e: 
        print("\n❌ 치명적 오류 발생") 
        print(e) 
        input("\n엔터 누르면 종료...")